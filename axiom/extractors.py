from __future__ import annotations

import json
import re
import os
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree


class MissingExtractor(RuntimeError):
    pass


@dataclass(frozen=True)
class ExtractedSegment:
    text: str
    modality: str
    page_number: int | None = None
    start_timestamp: str | None = None
    end_timestamp: str | None = None


TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".log"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
DOC_EXTENSIONS = {".doc"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
AUDIO_EXTENSIONS = {".wav", ".mp3", ".m4a", ".flac", ".ogg"}
SIDECAR_SUFFIXES = (
    ".ocr.txt",
    ".caption.txt",
    ".transcript.txt",
    ".transcript.json",
    ".pdf.txt",
    ".doc.txt",
    ".docx.txt",
    ".pdf.extracted.txt",
    ".doc.extracted.txt",
    ".docx.extracted.txt",
)


def file_type_for(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in PDF_EXTENSIONS:
        return "pdf"
    if suffix in DOCX_EXTENSIONS:
        return "docx"
    if suffix in DOC_EXTENSIONS:
        return "doc"
    if suffix in IMAGE_EXTENSIONS:
        return "image"
    if suffix in AUDIO_EXTENSIONS:
        return "audio"
    if suffix in TEXT_EXTENSIONS:
        return "text"
    return "unsupported"


def is_sidecar(path: Path) -> bool:
    name = path.name.lower()
    return any(name.endswith(suffix) for suffix in SIDECAR_SUFFIXES)


def extract_segments(path: Path) -> list[ExtractedSegment]:
    kind = file_type_for(path)
    if kind == "text":
        return [ExtractedSegment(path.read_text(encoding="utf-8", errors="replace"), "text")]
    if kind == "pdf":
        return extract_pdf(path)
    if kind == "docx":
        return extract_docx(path)
    if kind == "doc":
        return extract_doc(path)
    if kind == "image":
        return extract_image(path)
    if kind == "audio":
        return extract_audio(path)
    raise MissingExtractor(f"Unsupported file type: {path.suffix or '<none>'}")


def extract_pdf(path: Path) -> list[ExtractedSegment]:
    sidecar_text = read_extracted_text_sidecar(path)
    if sidecar_text:
        return [ExtractedSegment(sidecar_text, "text")]

    try:
        import fitz  # type: ignore
    except ImportError:
        return extract_pdf_pypdf(path)
    try:
        segments: list[ExtractedSegment] = []
        with fitz.open(path) as document:
            for index, page in enumerate(document, start=1):
                text = page.get_text("text").strip()
                if text:
                    segments.append(ExtractedSegment(text=text, modality="text", page_number=index))
        return segments
    except Exception as exc:  # noqa: BLE001 - adapter should report local extraction failure.
        fallback = extract_pdf_pypdf(path, raise_on_missing=False)
        if fallback:
            return fallback
        raise MissingExtractor(f"PDF extraction failed locally: {exc}") from exc


def extract_pdf_pypdf(path: Path, *, raise_on_missing: bool = True) -> list[ExtractedSegment]:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:
        if raise_on_missing:
            raise MissingExtractor("PDF extraction requires local PyMuPDF or pypdf installation.") from exc
        return []

    reader = PdfReader(str(path))
    segments: list[ExtractedSegment] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            segments.append(ExtractedSegment(text=text, modality="text", page_number=index))
    return segments


def extract_docx(path: Path) -> list[ExtractedSegment]:
    sidecar_text = read_extracted_text_sidecar(path)
    if sidecar_text:
        return [ExtractedSegment(sidecar_text, "text")]

    try:
        from docx import Document  # type: ignore
    except ImportError:
        return extract_docx_xml(path)

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return [ExtractedSegment(text="\n".join(parts), modality="text")]


def extract_docx_xml(path: Path) -> list[ExtractedSegment]:
    try:
        with zipfile.ZipFile(path) as package:
            names = [name for name in package.namelist() if name.startswith("word/") and name.endswith(".xml")]
            xml_blobs = [package.read(name) for name in names if name in {"word/document.xml"} or name.startswith("word/header") or name.startswith("word/footer")]
    except (OSError, zipfile.BadZipFile) as exc:
        raise MissingExtractor("DOCX extraction requires python-docx or a valid DOCX zip package.") from exc

    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    parts: list[str] = []
    for blob in xml_blobs:
        try:
            root = ElementTree.fromstring(blob)
        except ElementTree.ParseError:
            continue
        for paragraph in root.findall(".//w:p", namespace):
            texts = [node.text or "" for node in paragraph.findall(".//w:t", namespace)]
            line = "".join(texts).strip()
            if line:
                parts.append(line)
    if not parts:
        raise MissingExtractor("DOCX extraction produced no text.")
    return [ExtractedSegment(text="\n".join(parts), modality="text")]


def extract_doc(path: Path) -> list[ExtractedSegment]:
    sidecar_text = read_extracted_text_sidecar(path)
    if sidecar_text:
        return [ExtractedSegment(sidecar_text, "text")]

    for command in doc_text_commands(path):
        try:
            completed = subprocess.run(command, capture_output=True, text=True, timeout=90)
        except (OSError, subprocess.TimeoutExpired):
            continue
        text = (completed.stdout or "").strip()
        if completed.returncode == 0 and text:
            return [ExtractedSegment(text=text, modality="text")]

    libreoffice_text = extract_doc_with_libreoffice(path)
    if libreoffice_text:
        return [ExtractedSegment(text=libreoffice_text, modality="text")]

    raise MissingExtractor(
        "Legacy DOC extraction requires a staged local adapter such as antiword, catdoc, LibreOffice, or textract."
    )


def doc_text_commands(path: Path) -> list[list[str]]:
    commands: list[list[str]] = []
    if shutil.which("antiword"):
        commands.append(["antiword", str(path)])
    if shutil.which("catdoc"):
        commands.append(["catdoc", str(path)])
    if os.environ.get("AXIOM_TEXTRACT_DOC"):
        commands.append([os.environ["AXIOM_TEXTRACT_DOC"], str(path)])
    return commands


def extract_doc_with_libreoffice(path: Path) -> str:
    binary = shutil.which("soffice") or shutil.which("libreoffice")
    if not binary:
        return ""
    with tempfile.TemporaryDirectory(prefix="axiom-doc-") as tmp:
        completed = subprocess.run(
            [binary, "--headless", "--convert-to", "txt:Text", "--outdir", tmp, str(path)],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if completed.returncode != 0:
            return ""
        candidates = sorted(Path(tmp).glob("*.txt"))
        for candidate in candidates:
            text = candidate.read_text(encoding="utf-8", errors="replace").strip()
            if text:
                return text
    return ""


def read_extracted_text_sidecar(path: Path) -> str:
    for candidate in (Path(str(path) + ".txt"), Path(str(path) + ".extracted.txt")):
        if candidate.exists():
            text = candidate.read_text(encoding="utf-8", errors="replace").strip()
            if text:
                return text
    return ""


def extract_image(path: Path) -> list[ExtractedSegment]:
    segments: list[ExtractedSegment] = []
    for sidecar, modality in sidecar_candidates(path, [".ocr.txt", ".caption.txt"]):
        if sidecar.exists():
            text = sidecar.read_text(encoding="utf-8", errors="replace").strip()
            if text:
                segments.append(ExtractedSegment(text=text, modality="ocr" if "ocr" in modality else "image_caption"))

    if segments:
        return segments

    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore
    except ImportError as exc:
        raise MissingExtractor(
            "Image extraction requires OCR sidecars or local pillow+pytesseract installation."
        ) from exc

    text = pytesseract.image_to_string(Image.open(path)).strip()
    if not text:
        raise MissingExtractor("OCR produced no text for image.")
    return [ExtractedSegment(text=text, modality="ocr")]


def extract_audio(path: Path) -> list[ExtractedSegment]:
    for sidecar, _ in sidecar_candidates(path, [".transcript.json", ".transcript.txt"]):
        if not sidecar.exists():
            continue
        if sidecar.name.lower().endswith(".json"):
            return parse_transcript_json(sidecar)
        text = sidecar.read_text(encoding="utf-8", errors="replace")
        if looks_like_srt(text):
            return parse_srt(text)
        return [ExtractedSegment(text=text.strip(), modality="transcript")]

    whisper_segments = extract_audio_with_whisper(path)
    if whisper_segments:
        return whisper_segments

    raise MissingExtractor(
        "Audio extraction requires a transcript sidecar or a configured local Whisper.cpp adapter."
    )


def extract_audio_with_whisper(path: Path) -> list[ExtractedSegment]:
    command = whisper_command(path)
    if not command:
        return []
    try:
        completed = subprocess.run(command, capture_output=True, text=True, timeout=900)
    except (OSError, subprocess.TimeoutExpired):
        return []
    text = (completed.stdout or "").strip()
    if completed.returncode == 0 and text:
        return [ExtractedSegment(text=clean_whisper_stdout(text), modality="transcript")]
    transcript = Path(str(path) + ".transcript.txt")
    if transcript.exists():
        text = transcript.read_text(encoding="utf-8", errors="replace").strip()
        if text:
            return [ExtractedSegment(text=text, modality="transcript")]
    return []


def whisper_command(path: Path) -> list[str]:
    configured = os.environ.get("AXIOM_WHISPER_CPP")
    if configured:
        return [configured, "-f", str(path), "-otxt"]
    binary = shutil.which("whisper-cli") or shutil.which("whisper")
    if binary:
        return [binary, "-f", str(path), "-otxt"]
    return []


def clean_whisper_stdout(text: str) -> str:
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped and not stripped.lower().startswith(("whisper", "system_info", "main:")):
            lines.append(stripped)
    return "\n".join(lines)


def sidecar_candidates(path: Path, suffixes: list[str]) -> list[tuple[Path, str]]:
    candidates: list[tuple[Path, str]] = []
    for suffix in suffixes:
        candidates.append((Path(str(path) + suffix), suffix))
        candidates.append((path.with_suffix(suffix), suffix))
    return candidates


def parse_transcript_json(path: Path) -> list[ExtractedSegment]:
    data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    if isinstance(data, dict) and isinstance(data.get("segments"), list):
        segments = data["segments"]
    elif isinstance(data, list):
        segments = data
    else:
        text = data.get("text", "") if isinstance(data, dict) else ""
        return [ExtractedSegment(text=text.strip(), modality="transcript")] if text.strip() else []

    extracted: list[ExtractedSegment] = []
    for item in segments:
        text = str(item.get("text", "")).strip()
        if not text:
            continue
        extracted.append(
            ExtractedSegment(
                text=text,
                modality="transcript",
                start_timestamp=format_timestamp(item.get("start")),
                end_timestamp=format_timestamp(item.get("end")),
            )
        )
    return extracted


def looks_like_srt(text: str) -> bool:
    return bool(re.search(r"\d{2}:\d{2}:\d{2},\d{3}\s+-->\s+\d{2}:\d{2}:\d{2},\d{3}", text))


def parse_srt(text: str) -> list[ExtractedSegment]:
    blocks = re.split(r"\n\s*\n", text.strip())
    segments: list[ExtractedSegment] = []
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if len(lines) < 3:
            continue
        timing = next((line for line in lines if "-->" in line), "")
        if not timing:
            continue
        start, end = [part.strip().replace(",", ".") for part in timing.split("-->", 1)]
        body = " ".join(line for line in lines if line != timing and not line.isdigit())
        if body:
            segments.append(ExtractedSegment(body, "transcript", start_timestamp=start, end_timestamp=end))
    return segments


def format_timestamp(value: object) -> str | None:
    if value is None:
        return None
    if isinstance(value, str):
        return value
    try:
        seconds = float(value)
    except (TypeError, ValueError):
        return str(value)
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    return f"{hours:02d}:{minutes:02d}:{secs:06.3f}"
